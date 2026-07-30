import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from ebooklib import epub
import markdown

class BookExporter:
    def __init__(self, project, content):
        self.project = project
        self.content = content
        self.output_dir = Path(__file__).parent.parent.parent / "exports"
        self.output_dir.mkdir(exist_ok=True)
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.base_name = f"{project.name.replace(' ', '_')}_{self.timestamp}"

    def export_docx(self):
        """ייצוא ל-DOCX (Microsoft Word)"""
        doc = Document()
        
        title = doc.add_heading(self.project.name, 0)
        title.alignment = 1
        
        doc.add_paragraph(f"נוצר: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"סוג: {self.project.book_type}")
        doc.add_paragraph(f"מילים: {len(self.content.split())}")
        doc.add_paragraph("=" * 50)
        
        lines = self.content.split('\n')
        for line in lines:
            if line.startswith('#'):
                level = min(line.count('#'), 4)
                doc.add_heading(line.strip('#').strip(), level)
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        file_path = self.output_dir / f"{self.base_name}.docx"
        doc.save(file_path)
        print(f"   📄 DOCX: {file_path}")
        return file_path

    def export_epub(self):
        """ייצוא ל-EPUB (קוראים דיגיטליים)"""
        book = epub.EpubBook()
        
        book.set_identifier(self.project.id)
        book.set_title(self.project.name)
        book.set_language('he')
        book.add_author('AI Production Engine')
        
        content_html = markdown.markdown(self.content)
        html_content = f"""
        <html>
        <head><style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }}
            h1 {{ color: #6366f1; }}
        </style></head>
        <body>
            <h1>{self.project.name}</h1>
            <p><small>נוצר: {datetime.now().strftime('%d/%m/%Y %H:%M')}</small></p>
            <hr>
            {content_html}
        </body>
        </html>
        """
        
        chapter = epub.EpubHtml(
            title=self.project.name,
            file_name='chapter_1.xhtml',
            lang='he'
        )
        chapter.content = html_content
        book.add_item(chapter)
        
        book.toc = [(epub.Section(''), (chapter,))]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        style = '''
        @namespace epub "http://www.idpf.org/2007/ops";
        body { font-family: Arial, sans-serif; line-height: 1.6; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        file_path = self.output_dir / f"{self.base_name}.epub"
        epub.write_epub(file_path, book, {})
        print(f"   📖 EPUB: {file_path}")
        return file_path

    def export_pdf(self):
        """ייצוא ל-PDF באמצעות pdfkit"""
        try:
            import pdfkit
            config = pdfkit.configuration(wkhtmltopdf=r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe')
            
            content_html = markdown.markdown(self.content)
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>{self.project.name}</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 40px; max-width: 800px; margin: auto; }}
                    h1 {{ color: #6366f1; border-bottom: 2px solid #2a2a3a; padding-bottom: 10px; }}
                    h2 {{ color: #a855f7; margin-top: 30px; }}
                    h3 {{ color: #818cf8; }}
                    .meta {{ color: #888; font-size: 0.9em; margin-bottom: 30px; }}
                    hr {{ border: 1px solid #2a2a3a; margin: 30px 0; }}
                </style>
            </head>
            <body>
                <h1>{self.project.name}</h1>
                <div class="meta">
                    <p>📚 סוג: {self.project.book_type}</p>
                    <p>📝 מילים: {len(self.content.split())}</p>
                    <p>📅 תאריך: {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
                </div>
                <hr>
                {content_html}
            </body>
            </html>
            """
            
            file_path = self.output_dir / f"{self.base_name}.pdf"
            pdfkit.from_string(html_content, str(file_path), configuration=config)
            print(f"   📕 PDF: {file_path}")
            return file_path
        except Exception as e:
            print(f"   ⚠️ PDF נכשל: {e}")
            raise

    def export_all(self):
        """ייצוא לכל הפורמטים"""
        print("\n📦 מייצא ספר...")
        results = {}
        results['docx'] = self.export_docx()
        results['epub'] = self.export_epub()
        try:
            results['pdf'] = self.export_pdf()
        except Exception as e:
            print(f"   ⚠️ PDF לא נוצר: {e}")
        print(f"\n✅ כל הפורמטים נשמרו ב: {self.output_dir}")
        return results